import os
import re
from fastapi import FastAPI, Request
from fastapi.responses import HTMLResponse, FileResponse, JSONResponse
from fastapi.templating import Jinja2Templates
from pydantic import BaseModel
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.prompts import ChatPromptTemplate
from docx import Document

# Google Gemini API Key (hardcoded for now)
os.environ["GOOGLE_API_KEY"] = "AIzaSyBS6koInq5vWe8TRyYk0Azyq3Q-H0XkC2o"

app = FastAPI()
templates = Jinja2Templates(directory="templates")

# ------------------ HTML Routes ------------------ #
@app.get("/", response_class=HTMLResponse)
def home(request: Request):
    return templates.TemplateResponse("option.html", {"request": request})

@app.get("/createpaper", response_class=HTMLResponse)
def create_paper(request: Request):
    return templates.TemplateResponse("createpaper.html", {"request": request})

@app.get("/evaluateanswersheet", response_class=HTMLResponse)
def evaluate_sheet(request: Request):
    return templates.TemplateResponse("evaluateanswersheet.html", {"request": request})

# ------------------ API ------------------ #
class QuestionRequest(BaseModel):
    topic: str
    difficulty: int

@app.post("/generatequestions")
async def generate_questions(req: QuestionRequest):
    # Initialize Gemini LLM
    llm = ChatGoogleGenerativeAI(model="gemini-2.0-flash", temperature=0.7)

    # Prompt template
    prompt = ChatPromptTemplate.from_messages([
        ("system", "You are an expert exam paper generator."),
        ("user", """
Generate exam-style questions on the topic "{topic}".
Difficulty level: {difficulty} (1 = very easy, 10 = very hard).

Rules:
- If the user prompt explicitly specifies number of questions, types (MCQ, subjective, true/false, etc.), or formatting, follow those instructions.
- If the user does NOT specify, then default to:
    • 5 subjective questions
    • Number them as Q1, Q2, etc.
- For MCQ: Give 4 options (A–D) and bold the correct answer.
- For subjective: Keep them descriptive or analytical.
- For true/false: Phrase statements clearly.
- Ensure consistent formatting with clear separation.

Output format:
Q1. <question text>
[Options if applicable, one per line]
Answer (if applicable): <answer>

Q2. ...
""")
    ])

    # Run chain
    chain = prompt | llm
    response = await chain.ainvoke({
        "topic": req.topic,
        "difficulty": req.difficulty
    })

    # Extract and clean text
    questions_text = response.content.strip()
    questions_list = [line.strip() for line in questions_text.splitlines() if line.strip()]

    # Build Word document
    doc = Document()
    doc.add_heading(f"Question Paper - {req.topic} (Difficulty {req.difficulty})", 0)

    for q in questions_list:
        if q.startswith("Q"):
            doc.add_paragraph(q, style="List Number")  # numbered
        elif q.startswith(("A.", "B.", "C.", "D.")):
            doc.add_paragraph(q, style="List Bullet")  # options
        elif q.lower().startswith("answer"):
            p = doc.add_paragraph()
            run = p.add_run(q)
            run.bold = True
        else:
            doc.add_paragraph(q)

    # Sanitize topic for safe filename
    safe_topic = re.sub(r'[^a-zA-Z0-9_-]', '_', req.topic)
    filename = f"question_paper_{safe_topic}.docx"

    # Save file
    save_dir = "generated_files"
    os.makedirs(save_dir, exist_ok=True)
    filepath = os.path.join(save_dir, filename)
    doc.save(filepath)

    return JSONResponse({
        "questions": questions_list,
        "download_url": f"/download/{filename}"
    })

@app.get("/download/{filename}")
async def download_file(filename: str):
    filepath = os.path.join("generated_files", filename)
    if not os.path.exists(filepath):
        return JSONResponse({"error": "File not found"}, status_code=404)

    return FileResponse(
        filepath,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename
    )

import os
import re
from fastapi import UploadFile, File, Form
from fastapi.responses import JSONResponse
from docx import Document
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.prompts import ChatPromptTemplate

@app.post("/evaluate")
async def evaluate_answer_sheet(
    questionFile: UploadFile = File(...),
    answerFile: UploadFile = File(...),
    strictness: int = Form(5)
):
    try:
        # Read uploaded Question Paper
        q_doc = Document(questionFile.file)
        questions_text = "\n".join([p.text for p in q_doc.paragraphs if p.text.strip()])

        # Read uploaded Answer Sheet
        a_doc = Document(answerFile.file)
        answers_text = "\n".join([p.text for p in a_doc.paragraphs if p.text.strip()])

        # Initialize Gemini LLM
        llm = ChatGoogleGenerativeAI(model="gemini-2.0-flash", temperature=0.3)

        # Prompt Template
        prompt = ChatPromptTemplate.from_messages([
            ("system", "You are an expert examiner evaluating answer sheets."),
            ("user", """
Evaluate the student's answer sheet against the given question paper.

Strictness level: {strictness} (1 = very lenient, 10 = very strict).

Question Paper:
{questions}

Answer Sheet:
{answers}

Rules:
- Provide per-question evaluation (marks, comments, correctness).
- Give a total score at the end.
- Be consistent with strictness.
- Highlight strengths and weaknesses.
- Keep the feedback concise but clear.
- Output in structured text.

Format:
Q1: <evaluation + marks>
Q2: ...
...
Total: <marks>/<max_marks>
Comments: <overall feedback>
""")
        ])

        # Run LLM Chain
        chain = prompt | llm
        response = await chain.ainvoke({
            "questions": questions_text,
            "answers": answers_text,
            "strictness": strictness
        })

        evaluation_result = response.content.strip()

        return JSONResponse({"result": evaluation_result})

    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)
